import json
import re
from functools import reduce
from typing import Iterable, Generator

import requests
from bs4 import BeautifulSoup, Tag
from docx import Document

DOCUMENT_NAME = "acatistier.docx"
PRAYER_KEYWORDS = ("Psalmul 142", "Psalmul 50", "Condac", "Condacul", "Icos", "Icosul")
REPLACEMENTS = {
    "<br/>": "\n\t",
    "\x02": "",
}
TAG_PATTERN = re.compile('<.*?>')


def prepare_soups(urls: Iterable[str]) -> Generator[BeautifulSoup, None, None]:
    for url in urls:
        response = requests.get(url)
        response.raise_for_status()
        yield BeautifulSoup(response.text, "html.parser")


def add_summary(document: Document, summary: list[str]) -> None:
    document.add_heading("Acatistier - Cuprins", 0)
    document.add_paragraph("\n".join(summary))
    document.add_page_break()


def clean_html_text(text: Tag) -> str:
    text = reduce(
        lambda string, substitutions: string.replace(*substitutions),
        REPLACEMENTS.items(),
        text.decode_contents(),
    )
    return re.sub(TAG_PATTERN, '', text)


def add_prayers(document: Document, soups: Iterable[BeautifulSoup]):
    for i, soup in enumerate(soups):
        # Add title
        title = soup.select_one("#pagetitle").text.strip()
        print(f" -> {i}. Load \"{title}\"")
        document.add_heading(title, level=0)

        # Add content
        for paragraph in soup.select(".field-item p")[1:]:  # Skip first that just mentions starters
            # Match headings
            paragraph_text = paragraph.text.strip()
            if paragraph_text.startswith(PRAYER_KEYWORDS):
                document.add_heading(paragraph_text, level=1)
            # Match paragraphs
            else:
                document.add_paragraph(clean_html_text(paragraph))

        # Add break after content
        document.add_page_break()


def download(soups: None | list[BeautifulSoup] = None) -> None:
    """
    Retrieve saved catalog info, get content for each item and save document.

    :param soups: Provide only if you want to reuse cached soups to avoid unnecessary requests.
        Helpful for `playground` jupyter notebook.
    """
    with open("catalog.json", "r") as f:
        content = json.load(f)
    urls = content.keys()
    summary = content.values()

    if soups is None:
        print(f"⚠️ No soups provided. Ordering {len(content)} soups...")
        soups = prepare_soups(urls)

    document = Document()
    add_summary(document, summary)
    add_prayers(document, soups)
    document.save(DOCUMENT_NAME)


if __name__ == '__main__':
    download()
